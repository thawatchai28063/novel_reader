from __future__ import annotations

import json
import re
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


API_BASE = "http://localhost/novel_api/index.php"
NOVEL_ID = 3
OUTPUT = Path(r"C:\Users\it-ae\Documents\Codex\2026-05-24\files-mentioned-by-the-user-1\exports\เจ้าของร้านพิศวง_แก้คำ.docx")
COVER = Path(r"C:\Users\it-ae\Documents\Codex\2026-05-24\files-mentioned-by-the-user-1\novel_api\covers\owner_store_mystery.jpg")


def api_get(action: str, **params: str | int) -> object:
    query = {"action": action, **{key: str(value) for key, value in params.items()}}
    url = API_BASE + "?" + urllib.parse.urlencode(query)
    with urllib.request.urlopen(url, timeout=30) as response:
        payload = json.loads(response.read().decode("utf-8"))
    if not payload.get("ok"):
        raise RuntimeError(payload.get("error", "API error"))
    return payload.get("data")


def count_words(text: str) -> int:
    return len(re.findall(r"[\w\u0E00-\u0E7F]+", text, flags=re.UNICODE))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, name: str, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_editing_note(document: Document) -> None:
    table = document.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    rows = [
        ("วิธีใช้", "แก้คำผิดในเนื้อหาได้โดยตรง ห้ามลบหัวตอนที่ขึ้นต้นด้วย \"ตอนที่\" ถ้าต้องการนำกลับเข้าฐานข้อมูล"),
        ("ฐานข้อมูล", f"novel_id = {NOVEL_ID}"),
        ("แหล่งข้อมูล", "ข้อมูลนี้ export จาก MySQL ผ่าน PHP API ของแอป"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        row.cells[0].text = label
        row.cells[1].text = value
    document.add_paragraph()


def build_docx() -> None:
    novels = api_get("novels")
    novel = next(item for item in novels if int(item["id"]) == NOVEL_ID)
    chapters = api_get("chapters", novel_id=NOVEL_ID)

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Leelawadee UI"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(8)

    title_style = styles["Title"]
    title_style.font.name = "Leelawadee UI"
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
    title_style.font.size = Pt(26)
    title_style.font.color.rgb = RGBColor.from_string("1F3A35")

    heading = styles["Heading 1"]
    heading.font.name = "Leelawadee UI"
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), "Leelawadee UI")
    heading.font.size = Pt(18)
    heading.font.color.rgb = RGBColor.from_string("496B5B")
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(10)

    title = document.add_paragraph(style="Title")
    title_run = title.add_run(str(novel["title"]))
    set_font(title_run, "Leelawadee UI", 26, bold=True, color="1F3A35")

    subtitle = document.add_paragraph()
    run = subtitle.add_run(f"ไฟล์สำหรับแก้คำ | {len(chapters)} ตอน")
    set_font(run, "Leelawadee UI", 13, color="555555")

    if COVER.is_file():
        document.add_picture(str(COVER), width=Inches(2.0))

    add_editing_note(document)

    for index, summary in enumerate(chapters):
        if index > 0:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        chapter = api_get("chapter", id=int(summary["id"]))
        chapter_no = int(chapter["chapter_no"])
        chapter_title = str(chapter["title"]).strip() or f"ตอนที่ {chapter_no}"
        head = document.add_paragraph(style="Heading 1")
        head.add_run(f"ตอนที่ {chapter_no}: {chapter_title}")

        meta = document.add_paragraph()
        meta_run = meta.add_run(f"คำประมาณ {count_words(str(chapter['content']))} คำ")
        set_font(meta_run, "Leelawadee UI", 10, color="777777")

        for block in re.split(r"\n\s*\n", str(chapter["content"]).strip()):
            block = block.strip()
            if not block:
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Inches(0.28)
            paragraph.add_run(block)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)
    print(f"chapters={len(chapters)}")


if __name__ == "__main__":
    build_docx()
